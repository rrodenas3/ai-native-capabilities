"""Document ingestion pipeline for Cap-01 Decision Intelligence."""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Literal

from core.memory import Document, SemanticMemory
from core.utils.settings import get_settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".html", ".htm"}
AccessTier = Literal["public", "internal", "restricted"]


@dataclass(slots=True)
class IngestMetadata:
    title: str
    author: str | None = None
    date: str | None = None
    doc_type: str = "document"
    access_tier: AccessTier = "internal"
    source: str | None = None

    def as_document_metadata(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "author": self.author,
            "date": self.date,
            "doc_type": self.doc_type,
            "access_tier": self.access_tier,
            "source": self.source,
            "embedding_model": get_settings().LLM_EMBEDDINGS,
        }


class DocumentIngestor:
    """Load supported files and index them through SemanticMemory."""

    def __init__(self, semantic_memory: SemanticMemory) -> None:
        self.semantic_memory = semantic_memory

    def ingest_file(
        self,
        path: str | Path,
        *,
        metadata: IngestMetadata | None = None,
        doc_id: str | None = None,
    ) -> str:
        file_path = Path(path)
        self._validate_supported(file_path)
        content = self._read_file(file_path)
        resolved_metadata = metadata or IngestMetadata(
            title=file_path.stem,
            doc_type=file_path.suffix.lstrip(".").lower(),
            source=str(file_path),
        )
        document = Document(
            doc_id=doc_id or file_path.stem,
            capability="cap-01",
            content=content,
            metadata=resolved_metadata.as_document_metadata(),
            access_tier=resolved_metadata.access_tier,
        )
        self.semantic_memory.index([document])
        return document.doc_id

    def ingest_directory(
        self,
        path: str | Path,
        *,
        access_tier: AccessTier = "internal",
    ) -> list[str]:
        root = Path(path)
        files = sorted(file for file in root.rglob("*") if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS)
        documents: list[Document] = []
        doc_ids: list[str] = []
        for file_path in files:
            metadata = IngestMetadata(
                title=file_path.stem,
                doc_type=file_path.suffix.lstrip(".").lower(),
                access_tier=access_tier,
                source=str(file_path),
            )
            doc_id = file_path.stem
            doc_ids.append(doc_id)
            documents.append(
                Document(
                    doc_id=doc_id,
                    capability="cap-01",
                    content=self._read_file(file_path),
                    metadata=metadata.as_document_metadata(),
                    access_tier=access_tier,
                )
            )
        if documents:
            self.semantic_memory.index(documents)
        return doc_ids

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8")
        if suffix in {".html", ".htm"}:
            return _html_to_text(path.read_text(encoding="utf-8"))
        if suffix == ".pdf":
            return _read_pdf(path)
        if suffix == ".docx":
            return _read_docx(path)
        raise ValueError(f"Unsupported document format: {path.suffix}")

    @staticmethod
    def _validate_supported(path: Path) -> None:
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported document format: {path.suffix}")
        if not path.exists():
            raise FileNotFoundError(path)


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        stripped = data.strip()
        if stripped:
            self.parts.append(stripped)


def _html_to_text(html: str) -> str:
    parser = _TextHTMLParser()
    parser.feed(html)
    return " ".join(parser.parts)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx is required to ingest DOCX files") from exc

    document = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)

